"""Generate the accreditation-ready CO·PO document (.docx) from a scored matrix.

Unlike the old export (which omitted the mapping entirely), this produces the
NBA-standard CO·PO articulation matrix plus a per-cell CSAS justification appendix —
the audit trail that makes the mapping defensible.
"""
from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

POS = [f"PO{i}" for i in range(1, 13)]
BLOOM_NAME = {1: "Remember", 2: "Understand", 3: "Apply", 4: "Analyze", 5: "Evaluate", 6: "Create"}

# Cell fill by strength level (hex, no '#').
_LEVEL_FILL = {0: None, 1: "DBEAFE", 2: "93C5FD", 3: "3B82F6"}
_LEVEL_LABEL = {0: "None", 1: "Low", 2: "Medium", 3: "High"}
_ACCENT = RGBColor(0x4F, 0x46, 0xE5)
_MUTED = RGBColor(0x66, 0x66, 0x66)


def _shade_cell(cell, hex_fill: str | None) -> None:
    if not hex_fill:
        return
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_text(cell, text: str, *, bold=False, align_center=False, size=None, color=None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    if align_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def build_course_docx(meta: dict, matrix: list[dict]) -> bytes:
    """meta: {code,title,branch,semester}; matrix: rows of {co,bloom_level,pos,details}."""
    doc = Document()

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CO·PO MAPPING REPORT")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = _ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Course Alignment & Mapping Portal (CAMP 2.0) — CSAS deterministic engine")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = _MUTED

    # ---- Course information ----
    doc.add_paragraph()
    _heading(doc, "1. Course Information")
    info = doc.add_table(rows=5, cols=2)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Course Code", meta.get("code") or "—"),
        ("Course Title", meta.get("title") or "—"),
        ("Branch", meta.get("branch") or "—"),
        ("Semester", meta.get("semester") or "—"),
        ("Report Date", date.today().isoformat()),
    ]
    for i, (k, v) in enumerate(rows):
        _set_cell_text(info.rows[i].cells[0], k, bold=True)
        _set_cell_text(info.rows[i].cells[1], str(v))

    # ---- Course outcomes ----
    doc.add_paragraph()
    _heading(doc, "2. Course Outcomes")
    for i, row in enumerate(matrix, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"CO{i}: ")
        run.bold = True
        p.add_run(row["co"])
        b = row.get("bloom_level")
        if b:
            tag = p.add_run(f"   [Bloom {b} · {BLOOM_NAME.get(b, '')}]")
            tag.italic = True
            tag.font.size = Pt(9)
            tag.font.color.rgb = _ACCENT

    # ---- CO·PO articulation matrix ----
    doc.add_paragraph()
    _heading(doc, "3. CO·PO Articulation Matrix")
    table = doc.add_table(rows=1, cols=1 + len(POS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], "CO", bold=True, align_center=True)
    for j, po in enumerate(POS, 1):
        _set_cell_text(hdr[j], po, bold=True, align_center=True, size=8)

    totals = {po: 0 for po in POS}
    for i, row in enumerate(matrix, 1):
        cells = table.add_row().cells
        _set_cell_text(cells[0], f"CO{i}", bold=True, align_center=True)
        for j, po in enumerate(POS, 1):
            lvl = int(row["pos"].get(po, 0))
            totals[po] += lvl
            _set_cell_text(cells[j], "-" if lvl == 0 else str(lvl), align_center=True)
            _shade_cell(cells[j], _LEVEL_FILL[lvl])

    # average attainment row
    avg = table.add_row().cells
    _set_cell_text(avg[0], "Avg", bold=True, align_center=True)
    n = max(1, len(matrix))
    for j, po in enumerate(POS, 1):
        _set_cell_text(avg[j], f"{totals[po] / n:.1f}", align_center=True, size=8)

    _legend(doc)

    # ---- Justification appendix ----
    doc.add_paragraph()
    _heading(doc, "4. Mapping Justification (CSAS)")
    note = doc.add_paragraph()
    nr = note.add_run(
        "Each mapping below is computed deterministically by the CSAS engine from the "
        "CO's Bloom level, semantic similarity to the PO descriptor, and curated lexical "
        "affinity. Scores are reproducible and independent of any language model."
    )
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = _MUTED

    for i, row in enumerate(matrix, 1):
        mapped = [d for d in row.get("details", []) if d["level"] > 0]
        h = doc.add_paragraph()
        hr = h.add_run(f"CO{i}")
        hr.bold = True
        if not mapped:
            doc.add_paragraph("No significant PO mapping.").italic = True
            continue
        for d in sorted(mapped, key=lambda x: -x["level"]):
            p = doc.add_paragraph(style="List Bullet")
            lead = p.add_run(f"{d['po']} (Level {d['level']} · {_LEVEL_LABEL[d['level']]}): ")
            lead.bold = True
            p.add_run(d.get("rationale", ""))
            terms = d.get("matched_terms") or []
            if terms:
                t = p.add_run("  — terms: " + ", ".join(x["term"] for x in terms[:5]))
                t.font.size = Pt(9)
                t.font.color.rgb = _MUTED

    # ---- Footer ----
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run("Generated by CAMP 2.0 · deterministic · explainable · audit-ready")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = _MUTED

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = _ACCENT


def _legend(doc) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Strength:  1 = Low   2 = Medium   3 = High   ( - = no significant mapping )")
    r.font.size = Pt(8)
    r.font.color.rgb = _MUTED


def safe_filename(code: str | None, title: str | None) -> str:
    base = "_".join(filter(None, [code or "", title or "course"]))
    cleaned = "".join(c if c.isalnum() or c in "-_" else "_" for c in base).strip("_")
    return (cleaned or "course") + "_CO_PO_Report.docx"
