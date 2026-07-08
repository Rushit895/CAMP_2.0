"""DOCX export tests — build the document and read it back with python-docx to
assert the accreditation-critical content is present. Stdlib unittest.
"""
from __future__ import annotations

import io
import unittest
import zipfile

from docx import Document

from app.engine.csas import score_matrix
from app.services import export_service

COS = [
    "Design and develop a software system using modern engineering tools.",
    "Analyze the time complexity of algorithms.",
]


def _build() -> bytes:
    rows = [r.to_dict() for r in score_matrix(COS)]
    meta = {"code": "21CS42", "title": "Data Structures", "branch": "CSE", "semester": "4"}
    return export_service.build_course_docx(meta, rows)


class ExportTests(unittest.TestCase):
    def test_produces_valid_docx_zip(self):
        data = _build()
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(data)))  # .docx is a zip
        self.assertGreater(len(data), 2000)

    def test_document_contains_key_sections(self):
        doc = Document(io.BytesIO(_build()))
        text = "\n".join(p.text for p in doc.paragraphs)
        for expected in [
            "CO·PO MAPPING REPORT",
            "Course Information",
            "Course Outcomes",
            "CO·PO Articulation Matrix",
            "Mapping Justification",
            "21CS42",  # course code appears in the info table
        ]:
            self.assertIn(expected, text + _table_text(doc), f"missing: {expected}")

    def test_matrix_table_has_po_columns_and_rows(self):
        doc = Document(io.BytesIO(_build()))
        matrix_table = None
        for t in doc.tables:
            headers = [c.text for c in t.rows[0].cells]
            if "PO1" in headers and "PO12" in headers:
                matrix_table = t
                break
        self.assertIsNotNone(matrix_table, "CO·PO matrix table not found")
        self.assertEqual(len(matrix_table.rows[0].cells), 13)  # CO + PO1..12
        # header + 2 CO rows + average row
        self.assertEqual(len(matrix_table.rows), 4)
        # a design CO should show a strong PO3 somewhere in the body
        body_vals = [c.text for r in matrix_table.rows[1:] for c in r.cells]
        self.assertIn("3", body_vals)

    def test_filename_is_safe(self):
        name = export_service.safe_filename("21 CS/42", "Data Structures & Algo")
        self.assertTrue(name.endswith(".docx"))
        self.assertNotIn("/", name)
        self.assertNotIn(" ", name)


def _table_text(doc) -> str:
    out = []
    for t in doc.tables:
        for r in t.rows:
            out.extend(c.text for c in r.cells)
    return "\n".join(out)


if __name__ == "__main__":
    unittest.main(verbosity=2)
